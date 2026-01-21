import docx
import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fill_report(template_path, output_path):
    doc = docx.Document(template_path)

    # Content to insert
    project_title = "AI-Based Customer Churn Prediction & Action Recommendation System"
    
    problem_domain = (
        "In the rapidly evolving landscape of subscription-based business models—ranging from telecommunications and streaming services to SaaS (Software as a Service) platforms—customer retention has emerged as a critical metric for long-term sustainability. 'Customer Churn' refers to the phenomenon where customers discontinue their relationship with a service provider. "
        "Studies indicate that acquiring a new customer can cost five to twenty-five times more than retaining an existing one. Consequently, reducing churn by even a small percentage can lead to a significant increase in overall profitability.\n\n"
        "However, the core problem facing many organizations is the reactive nature of their retention strategies. Typically, businesses realize a customer is unhappy only after they have already cancelled their subscription. Traditional methods of analysis often fail to identify complex, non-linear patterns in user behavior that signal impending departure. "
        "Furthermore, even when a high-risk customer is identified, business teams often lack insight into the specific reasons driving that customer's dissatisfaction—whether it be pricing, lack of engagement, or technical issues—making it difficult to formulate effective personalized interventions.\n\n"
        "The objective of this project is to bridge this gap by developing an intelligent, proactive Customer Churn Prediction System. This system is designed not merely to flag at-risk customers with high accuracy but to serve as a comprehensive decision-support tool. "
        "It aims to solve the 'Black Box' problem of AI by providing clear, interpretable explanations for every prediction and coupling those insights with automated, actionable business recommendations to prevent churn before it happens."
    )

    proposed_treatment = (
        "To address the challenges outlined in the Problem Domain, we propose a sophisticated Artificial Intelligence solution that integrates Machine Learning for prediction, Explainable AI (XAI) for transparency, and a Rule-Based Expert System for decision support. "
        "The solution is architected as a cohesive desktop application that allows business analysts to interact seamlessly with advanced analytical models.\n\n"
        "1. Predictive Modeling Core:\n"
        "At the heart of the system lies a robust machine learning pipeline. We employ and compare multiple classification algorithms, specifically Logistic Regression and Random Forest Classifiers, to determine the most effective model for the dataset. "
        "The model analyzes ten key behavioral and demographic features—including 'Days Since Last Login', 'Payment Failures', 'Watch Time', 'Support Calls', and 'Tenure'—to calculate a precise probability score (0-100%) indicating the likelihood of a customer churning.\n\n"
        "2. Explainability Layer (SHAP Integration):\n"
        "A critical innovation in our treatment is the integration of SHapley Additive exPlanations (SHAP). In many AI deployments, trust is a barrier because users do not understand why a model makes a specific decision. "
        "Our system overcomes this by calculating the marginal contribution of each feature to the final prediction. For every customer, the system generates a personalized explanation (e.g., 'This customer is at high risk primarily because their payment failed twice and they haven't logged in for 45 days'). "
        "This empowers non-technical staff to understand the root causes of customer dissatisfaction.\n\n"
        "3. Action Recommendation Engine:\n"
        "Prediction is only useful if it leads to action. Our proposed treatment includes a logic-based recommendation engine that translates probability scores and risk factors into tangible business strategies. "
        "For instance, if a customer is flagged as 'High Risk' due to 'Price Sensitivity', the system automatically suggests offering a targeted discount. If the risk stems from 'Low Engagement', it suggests a content re-engagement campaign. "
        "This closes the loop between data analysis and business execution."
    )

    plan_of_work = (
        "The development of this project follows a structured System Development Life Cycle (SDLC) approach, divided into six distinct phases to ensure robustness and quality:\n\n"
        "Phase 1: Requirement Analysis & Data Simulation\n"
        "We began by defining the key factors that influence customer churn in a subscription context. Since real-world proprietary data is often inaccessible, we developed a sophisticated data generation script ('generate_data.py') to create a synthetic dataset of 1,000 customers. "
        "This dataset was engineered to contain realistic correlations (e.g., users with more payment failures are more likely to churn) to ensure the model learns meaningful patterns.\n\n"
        "Phase 2: Data Preprocessing & Exploratory Data Analysis (EDA)\n"
        "Raw data is rarely ready for modeling. This phase involved cleaning the data, handling missing values, and performing feature engineering. Categorical variables like 'Gender' and 'Subscription Type' were transformed using Label Encoding, "
        "and numerical features were normalized using Standard Scaling to ensure that algorithms like Logistic Regression perform optimally.\n\n"
        "Phase 3: Model Training & Evaluation\n"
        "We implemented a training pipeline to test multiple algorithms. The dataset was split into training (80%) and testing (20%) sets. We evaluated models based on Accuracy, Precision, Recall, and F1-Score. "
        "The best-performing model was then serialized (saved) using 'joblib' for later use in the application.\n\n"
        "Phase 4: Explainability Implementation\n"
        "We integrated the SHAP library to interpret the trained model. This involved creating an explainer object that could take a specific data point and decompose the prediction into feature importance values, distinguishing between factors that push risk up versus those that pull it down.\n\n"
        "Phase 5: User Interface (UI) Development\n"
        "To make the system accessible, we developed a modern Graphical User Interface (GUI) using Python's Tkinter framework. The UI was designed with a focus on User Experience (UX), featuring a sidebar navigation, clean forms for data entry, and visual indicators for risk levels.\n\n"
        "Phase 6: Integration & Testing\n"
        "The final phase involved connecting the backend ML models with the frontend UI. Rigorous testing was conducted to ensure that inputs from the UI were correctly processed, predictions were accurate, and the recommendations displayed were logically consistent with the risk factors."
    )

    software_specs = (
        "The project relies on a modern, open-source technology stack selected for its reliability, performance, and ease of deployment.\n\n"
        "Software Requirements:\n"
        "1. Programming Language: Python 3.8+\n"
        "   Chosen for its extensive ecosystem of data science and machine learning libraries.\n\n"
        "2. Core Libraries & Frameworks:\n"
        "   - Pandas & NumPy: Utilized for high-performance data manipulation, structure, and numerical analysis.\n"
        "   - Scikit-Learn: The primary machine learning library used for training algorithms (Logistic Regression, Random Forest), preprocessing data, and evaluating metrics.\n"
        "   - SHAP (SHapley Additive exPlanations): A game-theoretic approach to explain the output of the machine learning model.\n"
        "   - Joblib: Used for object serialization, allowing us to save the trained model and encoders to disk and load them instantly during application runtime.\n"
        "   - Tkinter: Python's standard GUI library, used to create the desktop application interface without requiring a web browser.\n\n"
        "3. Development Environment:\n"
        "   - Visual Studio Code or PyCharm as the Integrated Development Environment (IDE).\n"
        "   - Git & GitHub for version control and source code management.\n\n"
        "Hardware Requirements:\n"
        "To ensure smooth training and inference, the following minimum specifications are recommended:\n"
        "- Processor: Intel Core i5 or equivalent (Minimum 2.0 GHz)\n"
        "- RAM: 8 GB or higher (to handle dataset loading and model training in memory)\n"
        "- Storage: 500 MB free space (for environment, libraries, and datasets)\n"
        "- Display: 1366x768 resolution or higher for optimal UI viewing."
    )

    user_guide = (
        "This User Guide provides a comprehensive walkthrough of the AI-Based Customer Churn Prediction System. The application is divided into several modules, each accessible via the left-hand navigation sidebar.\n\n"
        "1. Dashboard Overview (Home Page)\n"
        "   - Upon launching the application ('python ui/app.py'), you are greeted by the Home Dashboard.\n"
        "   - This page serves as the central hub, providing a brief overview of the system's capabilities and quick access buttons to the main features (Predict, Upload, Reports).\n\n"
        "2. Individual Customer Prediction (Predict Page)\n"
        "   - Navigate to the 'Predict' tab to analyze a single customer.\n"
        "   - Input Data: Fill in the customer details manually (Age, Tenure, Monthly Charges, etc.) or click 'Load Sample Data' to auto-populate the form for testing.\n"
        "   - Analyze: Click 'Predict Churn'. The AI analyzes the data in real-time.\n"
        "   - Results: The system displays the Churn Probability (%), Risk Level (Low/Med/High), and SHAP-based explanations ('Why this prediction?').\n"
        "   - Action Plan: Review the tailored business recommendations at the bottom (e.g., 'Offer 20% Discount').\n\n"
        "3. Batch Data Processing (Upload Page)\n"
        "   - For analyzing multiple customers at once, go to the 'Upload' tab.\n"
        "   - Click 'Browse' to select a CSV file containing customer records.\n"
        "   - The system validates the file and processes all records simultaneously.\n"
        "   - Results are displayed in a tabular format, showing the predicted risk for every customer in the file.\n\n"
        "4. Data Analytics & Visualization (Charts Page)\n"
        "   - The 'Charts' tab provides a visual summary of the customer base.\n"
        "   - View dynamic graphs such as the 'Churn vs. Retention Distribution' (Pie Chart) or 'Risk Factor Analysis' (Bar Charts).\n"
        "   - These visualizations help management understand broader trends beyond individual predictions.\n\n"
        "5. Report Generation (Reports Page)\n"
        "   - Navigate to the 'Reports' tab to document findings.\n"
        "   - You can generate a comprehensive PDF or Text report summarizing recent predictions and overall system performance.\n"
        "   - This feature is essential for sharing insights with stakeholders who do not have access to the app.\n\n"
        "6. System Configuration (Settings Page)\n"
        "   - The 'Settings' tab allows you to customize the AI's behavior.\n"
        "   - Model Selection: You can switch the active machine learning model (e.g., from 'Logistic Regression' to 'Random Forest') to compare performance or use a more complex algorithm.\n"
        "   - Theme: Toggle between Dark and Light modes for better visibility."
    )

    # Strategy: Iterate through paragraphs and find headers to append text after
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # 1. Project Title
        if "Project Title:" in text and "___" in text:
            paragraph.text = text.split(":")[0] + ": " + project_title
            continue

        # 2. Problem Domain
        if "1. PROBLEM DOMAIN" in text:
            paragraph.add_run("\n" + problem_domain)

        # 3. Proposed Treatment
        if "2. PROPOSED TREATMENT" in text:
             paragraph.add_run("\n" + proposed_treatment)

        # 4. Plan of Work
        if "3. PLAN OF WORK" in text:
             paragraph.add_run("\n" + plan_of_work)
             
        # 5. Software Specs
        if "5. SOFTWARE AND HARDWARE SPECIFICATIONS" in text:
             paragraph.add_run("\n" + software_specs)

        # 6. User Guide
        if "8. USER GUIDE" in text:
             paragraph.add_run("\n" + user_guide)

    doc.save(output_path)
    print(f"Report filled and saved to: {output_path}")

if __name__ == "__main__":
    template = "AI Project Report Template .docx"
    output = "AI_Project_Report_Filled.docx"
    
    if os.path.exists(template):
        fill_report(template, output)
    else:
        print(f"Template not found: {template}")
